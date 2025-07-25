from flask import Flask, render_template, request, send_file, redirect, url_for
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re
from collections import Counter
import spacy
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image as ReportLabImage, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from zipfile import ZipFile
import io

app = Flask(__name__, template_folder='.')

# Load spaCy model for NLP processing
nlp = spacy.load("en_core_web_sm")

uploaded_data = {
    "blocks": [],
    "positive": "2",
    "negative": "0.25",
    "range_start": 1,
    "range_end": 9999,
    "base": None
}

# Enhanced patterns with NLP support
QUESTION_PATTERN = re.compile(r"""
    ^                       # Start of line
    (?:                     # Non-capturing group
        Q                   # Q prefix
        |                   # OR
        Question\s          # "Question" followed by space
        |                   # OR
        \d{1,3}\.          # 1-3 digits followed by dot
    )
    \s*                     # Optional whitespace
    (?=\w)                  # Lookahead for word character
""", re.VERBOSE | re.IGNORECASE)

# Enhanced option pattern to handle various formats
OPTION_PATTERN = re.compile(r"""
    ^\s*                     
    (                       
        \d{2,2}\.           # 1.
        |                     
        \d{2,2}\)           # 1)
        |
        [a-z]\)             # a)
        |
        [a-z]\.             # a.
        |
        \([a-z]\)           # (a)
        |
        \([A-Z]\)           # (A)
        |
        [ivx]{1,4}\.       # i.
        |
        [A-Z]\.             # A. (capital letter)
        |
        [A-Z]\)             # A) (capital letter)
    )
    \s+                     
""", re.VERBOSE | re.IGNORECASE)

# New pattern to detect question numbers in text
QUESTION_NUMBER_PATTERN = re.compile(r'\b(\d{2,3})\.\s')

@app.route('/')
def index():
    return render_template('index.html')

def extract_questions_from_pdf(pdf_data):
    pdf_doc = fitz.open(stream=pdf_data, filetype="pdf")
    questions = []
    current_question = None
    current_images = []
    page_question_count = {}
    prev_line_blank = True  # Track if previous line was blank

    for page_number in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_number)
        text = page.get_text()
        
        # Extract images for this page
        page_images = []
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = pdf_doc.extract_image(xref)
            page_images.append({
                "bytes": base_image["image"],
                "ext": base_image["ext"],
                "width": base_image["width"],
                "height": base_image["height"]
            })
        
        # Initialize question count for this page
        page_question_count[page_number] = 0
        
        # Split text into lines for processing
        lines = text.split('\n')
        for line in lines:
            stripped_line = line.strip()
            is_blank = (stripped_line == '')
            
            # Check if this is a new question
            if (prev_line_blank and 
                QUESTION_PATTERN.search(stripped_line) and 
                not current_question and 
                not OPTION_PATTERN.search(stripped_line)):
                
                # Save previous question if exists
                if current_question:
                    questions.append((current_question, current_images))
                    current_images = []
                    page_question_count[page_number] += 1
                
                # Start new question
                current_question = stripped_line
                # Add images found so far to new question
                current_images.extend(page_images)
                page_images = []  # Reset for next images
            elif current_question:
                # Use NLP to verify if this is actually a new question
                if prev_line_blank and QUESTION_PATTERN.search(stripped_line):
                    spacy_doc = nlp(stripped_line)
                    # Check if this looks like a complete question
                    if any(token.is_title for token in spacy_doc) and not OPTION_PATTERN.search(stripped_line):
                        # Save current question
                        questions.append((current_question, current_images))
                        current_images = []
                        page_question_count[page_number] += 1
                        
                        # Start new question
                        current_question = stripped_line
                        current_images.extend(page_images)
                        page_images = []
                        continue
                
                # Accumulate lines for current question
                current_question += '\n' + stripped_line
            
            prev_line_blank = is_blank
        
        # After processing page, add any remaining images to current question
        if current_question:
            current_images.extend(page_images)
            page_images = []
    
    # Add last question if exists
    if current_question:
        questions.append((current_question, current_images))
        page_question_count[page_number] += 1
    
    return questions, page_question_count

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def force_table_indent_and_widths(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Inches(0.2).pt)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    tbl.insert(0, tblPr)
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.85)

def parse_correct_answer(candidate_text, raw_options):
    """Enhanced correct answer parser that handles various formats"""
    # Remove prefix if present
    candidate_text = re.sub(r'^correct answer\s*:\s*', '', candidate_text, flags=re.IGNORECASE).strip()
    
    # Handle cases like "2. Schedule II"
    if re.match(r'^\d+\..+', candidate_text):
        # Extract the number before the dot
        num_match = re.search(r'^(\d+)\.', candidate_text)
        if num_match:
            return num_match.group(1)
    
    # Handle cases like "Article 97"
    if re.match(r'^Article\s+\d+', candidate_text, re.IGNORECASE):
        art_match = re.search(r'(\d+)', candidate_text)
        if art_match:
            return art_match.group(1)
    
    # Handle letter-number combinations like "B. Governor"
    letter_match = re.match(r'^([A-D])[\.\)]\s*(.*)', candidate_text)
    if letter_match:
        letter = letter_match.group(1)
        return str(ord(letter) - ord('A') + 1)
    
    # Handle multiple options like "A, B and C"
    if ',' in candidate_text or ' and ' in candidate_text:
        # Extract all option letters
        option_letters = re.findall(r'[A-D]', candidate_text)
        if option_letters:
            # Convert to numbers (A->1, B->2, etc.)
            option_numbers = [str(ord(letter) - ord('A') + 1) for letter in option_letters]
            return ",".join(option_numbers)
    
    # Handle numeric answers like "1" or "2"
    num_match = re.search(r'\b(\d+)\b', candidate_text)
    if num_match:
        return num_match.group(1)
    
    # Handle letter answers like "B" or "C"
    letter_match = re.search(r'\b([A-D])\b', candidate_text)
    if letter_match:
        letter = letter_match.group(1)
        return str(ord(letter) - ord('A') + 1)
    
    # Fallback: try to match by option text content
    candidate_clean = OPTION_PATTERN.sub('', candidate_text).strip()
    for idx, opt in enumerate(raw_options):
        opt_clean = OPTION_PATTERN.sub('', opt).strip()
        if candidate_clean == opt_clean:
            return str(idx + 1)
    
    return candidate_text  # Default return if no pattern matches

def process_question_block(block, positive, negative):
    block_text, images = block
    # Preserve original lines
    lines = [line for line in block_text.split("\n") if line.strip()]
    opts = []
    raw_options = []
    ans = ''
    sol_lines = []
    question_lines = []
    
    # Extract question number with enhanced pattern
    q_num = ""
    q_num_match = re.search(r'(?:Q|Question\s)?(\d{1,3})\.?', block_text, re.IGNORECASE)
    if q_num_match:
        q_num = "Q" + q_num_match.group(1)  # Standardize to Q-prefix

    capturing_question = True
    capturing_option_index = -1
    capturing_solution = False

    # Track if we're in a multi-question block
    question_number_detected = False
    new_question_detected = False

    for line in lines:
        # Check for embedded question numbers
        if capturing_question and not capturing_option_index and QUESTION_NUMBER_PATTERN.search(line):
            question_number_detected = True
            new_question_detected = True
        
        # Check for option patterns
        if OPTION_PATTERN.search(line) and not new_question_detected:
            capturing_question = False
            capturing_solution = False
            raw_options.append(line)
            opts.append(OPTION_PATTERN.sub("", line).strip())
            capturing_option_index = len(opts) - 1
            question_number_detected = False

        elif capturing_option_index != -1 and not line.lower().startswith(("correct answer", "solution")) and not new_question_detected:
            opts[capturing_option_index] += ' ' + line
            raw_options[-1] += ' ' + line

        elif re.search(r'^correct answer', line, re.IGNORECASE):
            ans = parse_correct_answer(line, raw_options)
            capturing_option_index = -1
            capturing_solution = False
            question_number_detected = False

        elif line.lower().startswith("solution"):
            sol_lines.append(line.split(":", 1)[-1].strip())
            capturing_solution = True
            capturing_option_index = -1
            question_number_detected = False

        elif capturing_solution:
            sol_lines.append(line.strip())

        elif capturing_question and not new_question_detected:
            # Remove question number prefix
            clean_line = re.sub(r'(?:Q|Question\s)?\d{1,3}\.?', '', line).strip()
            question_lines.append(clean_line)
        
        # Reset new question detection for next line
        new_question_detected = False

    if len(opts) > 4:
        # Find the first option that looks like a new question
        cut_index = None
        for i, opt in enumerate(opts):
            if QUESTION_NUMBER_PATTERN.search(opt):
                cut_index = i
                break
        
        if cut_index is not None:
            # Move extra options to question text
            extra_raw_opts = raw_options[:cut_index]
            core_opts = opts[cut_index:cut_index+4] if len(opts) >= cut_index+4 else opts[cut_index:]
            question_lines.extend(extra_raw_opts)
            final_options = core_opts
        else:
            # Fallback to last 4 options
            extra_raw_opts = raw_options[:-4]
            core_opts = opts[-4:]
            question_lines.extend(extra_raw_opts)
            final_options = core_opts

    q = " ".join(question_lines)

    # Format options with proper line breaks
    if " a)" in q.lower() and " b)" in q.lower():
        q = re.sub(r'\s([a-z]\)', r'\n\1', q, flags=re.IGNORECASE)
    elif "1." in q and "2." in q:
        q = re.sub(r'\s(\d{1,2}\.)', r'\n\1', q)
    elif "i." in q and "ii." in q:
        q = re.sub(r'\s([ivx]{1,4}\.)', r'\n\1', q)

    solution = " ".join(sol_lines).strip()

    return {
        "Question": q.strip(),
        "Type": "multiple_choice",
        "Options": final_options,
        "Answer": ans,
        "Solution": solution,
        "Positive Marks": positive,
        "Negative Marks": negative,
        "Images": images,
        "Question Number": q_num
    }

def generate_docx(questions, bold_question=False):
    """Generate DOCX document from processed questions"""
    document = Document()
    doc_stream = BytesIO()
    
    for data in questions:
        # Create a table with question details
        table = document.add_table(rows=10, cols=2)
        table.autofit = False
        force_table_indent_and_widths(table)
        set_table_borders(table)

        labels = ["Question", "Type", "Option", "Option", "Option", "Option",
                "Answer", "Solution", "Positive Marks", "Negative Marks"]
        values = [data["Question"], data["Type"]] + data["Options"][:4] + [
            data["Answer"], data["Solution"], data["Positive Marks"], data["Negative Marks"]]

        for i, (label, value) in enumerate(zip(labels, values)):
            row = table.rows[i]
            row.cells[0].text = label

            if label == "Question":
                cell = row.cells[1]
                # Clear existing content
                for paragraph in cell.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)

                # Add question text with optional bold formatting
                p = cell.add_paragraph()
                if bold_question:
                    run = p.add_run(value)
                    run.bold = True
                else:
                    p.add_run(value)

                # Insert images associated with this question
                for img in data.get("Images", []):
                    try:
                        p.add_run().add_break()  # Line break before image
                        run = p.add_run()
                        run.add_picture(BytesIO(img["bytes"]), width=Inches(2))
                    except Exception as e:
                        print(f"Error adding image to DOCX: {e}")

            else:
                row.cells[1].text = re.sub(r"\s*\n\s*", " ", value).strip()

        document.add_paragraph("") #creating a line break between two table
    
    # Save document
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

def generate_pdf(questions, bold_question=False):
    """Generate PDF document from processed questions"""
    pdf_stream = BytesIO()
    doc = SimpleDocTemplate(pdf_stream, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    if bold_question:
        bold_style = ParagraphStyle(
            'BoldQuestion',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            textColor=colors.red,
        )
    else:
        bold_style = styles['Normal']

    for data in questions:
        normal_style = styles['Normal']

        def format_text_with_linebreaks(text):
            # Add <br/> before list-like patterns
            patterns = [
                r"\s([A-Za-z]\.)",
                r"\s(\d{1,2}\.)",
                r"\s([ivxlcdm]{1,4}\.)",
                r"\s([IVXLCDM]{1,4}\.)"
            ]

            for pattern in patterns:
                text = re.sub(pattern, r"<br/>&nbsp;\1", text)
            return text

        # Create table data
        table_data = [
            ["Question", Paragraph(format_text_with_linebreaks(data["Question"]), bold_style)],
            ["Type", data["Type"]],
            ["Option A", Paragraph(data["Options"][0], normal_style)],
            ["Option B", Paragraph(data["Options"][1], normal_style)],
            ["Option C", Paragraph(data["Options"][2], normal_style)],
            ["Option D", Paragraph(data["Options"][3], normal_style)],
            ["Answer", data["Answer"]],
            ["Solution", Paragraph(format_text_with_linebreaks(data["Solution"]), normal_style)],
            ["Positive Marks", data["Positive Marks"]],
            ["Negative Marks", data["Negative Marks"]]
        ]

        # Create PDF table
        table = Table(table_data, colWidths=[1.5 * inch, 5 * inch])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))

        elements.append(table)

        # Add image after table (simulating line break)
        if data.get("Images"):
            try:
                img = data["Images"][0]  # First image only
                img_stream = BytesIO(img["bytes"])
                aspect_ratio = img["height"] / img["width"]
                img_width = 2 * inch
                img_height = img_width * aspect_ratio

                if img_height > 2 * inch:
                    img_height = 2 * inch
                    img_width = img_height / aspect_ratio

                elements.append(Spacer(1, 0.1 * inch))  # newline effect
                elements.append(ReportLabImage(img_stream, width=img_width, height=img_height))
            except Exception as e:
                print(f"Error rendering image after table: {e}")

        elements.append(Spacer(1, 0.3 * inch))  # space before next question

    # Build PDF
    doc.build(elements)
    pdf_stream.seek(0)
    return pdf_stream

@app.route('/upload', methods=['POST'])
def upload():
    pdf_file = request.files['pdf_file']
    uploaded_data["original_filename"] = pdf_file.filename.rsplit('.', 1)[0]  # remove extension
    uploaded_data["positive"] = request.form.get('positive', '2')
    uploaded_data["negative"] = request.form.get('negative', '0.25')

    # Get the bold_question value from the form
    bold_question = request.form.get('bold_question', 'no') == 'yes'
    
    # Store it in uploaded_data
    uploaded_data["bold_question"] = bold_question

    if request.form.get("generate_all") == "yes":
        uploaded_data["range_start"] = 1
        uploaded_data["range_end"] = 9999
    else:
        try:
            uploaded_data["range_start"] = int(request.form.get('range_start') or 1)
            uploaded_data["range_end"] = int(request.form.get('range_end') or 9999)
        except ValueError:
            return "❌ Invalid range input. Please enter valid numbers or check 'Generate all questions'.", 400

    blocks, page_question_count = extract_questions_from_pdf(pdf_file.read())
    uploaded_data["blocks"] = blocks

    errors = []
    base_numbers = []
    option_issues = []
    repeated_questions = []
    pattern = r'(?:Q|Question\s)?(\d{1,3})\.?' 
    multi_page_warnings = []

    # Generate multi-page warnings
    for page_num, count in page_question_count.items():
        if count > 1:
            multi_page_warnings.append(f"Page {page_num+1} has {count} questions. Images on this page are associated with the first question that appears there.")

    last_qnum = None
    for i, block in enumerate(blocks):
        # Get block text - blocks are tuples of (text, images)
        block_text = block[0] if isinstance(block, tuple) else block
        match = re.search(pattern, block_text, re.IGNORECASE)
        if match:
            try:
                num = int(match.group(1))
                base_numbers.append(num)
                
                # Sequence check
                if last_qnum is not None and num != last_qnum + 1:
                    errors.append(f"Issue at Q{num} (expected Q{last_qnum + 1})")
                last_qnum = num
            
                # Count options
                lines = block_text.splitlines()
                option_count = sum(1 for line in lines if OPTION_PATTERN.search(line))
                if option_count != 4:
                    option_issues.append(f"Q{num} has {option_count} options")
            except (ValueError, IndexError):
                # Skip invalid question numbers
                continue

    # Calculate actual start/end from the extracted numbers
    actual_start = min(base_numbers) if base_numbers else 0
    actual_end = max(base_numbers) if base_numbers else 0  

    # Repeated questions
    counts = Counter(base_numbers)
    repeated_questions = [f"Q{num}" for num, count in counts.items() if count > 1]

    uploaded_data["base"] = base_numbers[0] if base_numbers else 1

    # Filter for selected question range
    filtered_qnums = []
    questions_to_generate = 0
    for block in blocks:
        block_text = block[0] if isinstance(block, tuple) else block
        match = re.search(pattern, block_text, re.IGNORECASE)
        if match:
            try:
                q_num = int(match.group(1))
                if uploaded_data["range_start"] <= q_num <= uploaded_data["range_end"]:
                    filtered_qnums.append(q_num)
                    questions_to_generate += 1
            except (ValueError, IndexError):
                # Skip invalid question numbers
                continue

    gen_start = min(filtered_qnums) if filtered_qnums else uploaded_data["range_start"]
    gen_end = max(filtered_qnums) if filtered_qnums else uploaded_data["range_end"]

    # Ensure lists are not None
    errors = errors or []
    option_issues = option_issues or []
    repeated_questions = repeated_questions or []
    multi_page_warnings = multi_page_warnings or []

    return render_template("diagnose.html",
        total_qs=len(blocks),
        actual_start=actual_start,
        actual_end=actual_end,
        range_start=uploaded_data["range_start"],
        range_end=uploaded_data["range_end"],
        base=uploaded_data["base"],
        option_issues=option_issues,
        errors=errors,
        repeated_questions=repeated_questions,
        questions_to_generate=questions_to_generate,
        gen_start=gen_start,
        gen_end=gen_end,
        multi_page_warnings=multi_page_warnings
    )

@app.route('/generate', methods=['POST'])
def generate():
    confirm = request.form.get("confirm", "no")
    output_format = request.form.get("format", "docx")
    blocks = uploaded_data["blocks"]
    positive = uploaded_data["positive"]
    negative = uploaded_data["negative"]
    range_start = uploaded_data["range_start"]
    range_end = uploaded_data["range_end"]
    bold_question = uploaded_data["bold_question"]  # Get the bold setting

    if confirm == "no":
        return redirect(url_for("index"))

    pattern = r'(?:Q|Question\s)?(\d{1,3})\.?' 
    selected_blocks = []

    for block in blocks:
        block_text = block[0] if isinstance(block, tuple) else block
        match = re.search(pattern, block_text, re.IGNORECASE)
        if match:
            try:
                q_num = int(match.group(1))
                if range_start <= q_num <= range_end:
                    selected_blocks.append(block)
            except (ValueError, IndexError):
                # Skip invalid question numbers
                continue

    if not selected_blocks:
        return "No questions found in the selected range.", 400

    # Process all selected questions
    processed_questions = []
    for block in selected_blocks:
        data = process_question_block(block, positive, negative)
        processed_questions.append(data)

    # Get a clean filename from the uploaded PDF name
    base_name = re.sub(r'[\\/*?:"<>|]', "_", uploaded_data.get("original_filename", "Processed_MCQs"))
    docx_filename = f"Bulk_Uploader_of_{base_name}.docx"
    pdf_filename = f"Bulk_Uploader_of_{base_name}.pdf"
    zip_filename = f"Bulk_Uploader_of_{base_name}.zip"

    # Handle different output formats
    if output_format == "docx":
        docx_stream = generate_docx(processed_questions, bold_question)
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=docx_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    elif output_format == "pdf":
        pdf_stream = generate_pdf(processed_questions, bold_question)
        return send_file(
            pdf_stream,
            as_attachment=True,
            download_name=pdf_filename,
            mimetype="application/pdf"
        )

    elif output_format == "zip":
        # Create ZIP with both DOCX and PDF
        docx_stream = generate_docx(processed_questions)
        pdf_stream = generate_pdf(processed_questions)
        
        zip_stream = BytesIO()
        with ZipFile(zip_stream, 'w') as zipf:
            zipf.writestr(docx_filename, docx_stream.getvalue())
            zipf.writestr(pdf_filename, pdf_stream.getvalue())
        
        zip_stream.seek(0)
        return send_file(
            zip_stream,
            as_attachment=True,
            download_name=zip_filename,
            mimetype="application/zip"
        )

    return "❌ Only DOCX, PDF, and ZIP formats are supported on this server.", 400

if __name__ == "__main__":
    app.run(host="0.0.0.0", debug=True)