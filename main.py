from flask import Flask, render_template, request, send_file, redirect, url_for
import fitz # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re
from collections import Counter
import platform
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image as ReportLabImage, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
import os
from zipfile import ZipFile
import io
from PIL import Image
import base64

app = Flask(__name__, template_folder='.')

uploaded_data = {
    "blocks": [],
    "positive": "2",
    "negative": "0.25",
    "range_start": 1,
    "range_end": 9999,
    "base": None
}

# ✅ Universal option pattern (allows leading whitespace)
OPTION_LABEL_RE = re.compile(r"^\s*[\(\[]?(\d{1,2}|[A-Za-z]|[ivxlcdmIVXLCDM]{1,4})[\)\.\]]\s*")

@app.route('/')
def index():
    return render_template('index.html')

def get_full_paragraph_text(para, list_counter):
    text = para.text.strip()
    p = para._element
    is_list_item = p.xpath(".//w:numPr")
    
    if is_list_item:
        level = p.xpath(".//w:ilvl")
        
        try:
            level_num = int(level[0].text)
        except (IndexError, TypeError, ValueError):
            level_num = 0

        # Maintain a list count per level
        if level_num not in list_counter:
            list_counter[level_num] = 1
        else:
            list_counter[level_num] += 1

        num_label = f"{list_counter[level_num]}. "
        return f"{num_label}{text}".strip()
    
    return text

def extract_images_from_docx_paragraphs(doc):
    """Enhanced image extraction from DOCX document with better error handling"""
    images_data = []
    
    print(f"Starting image extraction from DOCX document with {len(doc.paragraphs)} paragraphs")
    
    # Method 1: Extract from paragraphs (inline images) with context
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        
        # Check if paragraph contains images
        for run in para.runs:
            # Look for drawing elements (newer format)
            for drawing in run._element.xpath('.//w:drawing'):
                try:
                    # Get image relationship ID
                    blip = drawing.xpath('.//a:blip/@r:embed', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    })
                    
                    if blip:
                        rel_id = blip[0]
                        # Get the image from document relationships
                        if rel_id in doc.part.related_parts:
                            image_part = doc.part.related_parts[rel_id]
                            image_bytes = image_part.blob
                            
                            # Try to determine image format
                            image_format = 'png'  # default
                            if image_bytes.startswith(b'\xFF\xD8\xFF'):
                                image_format = 'jpeg'
                            elif image_bytes.startswith(b'\x89PNG'):
                                image_format = 'png'
                            elif image_bytes.startswith(b'GIF'):
                                image_format = 'gif'
                            elif image_bytes.startswith(b'BM'):
                                image_format = 'bmp'
                            
                            # Get image dimensions using PIL with better error handling
                            try:
                                with Image.open(BytesIO(image_bytes)) as img:
                                    width, height = img.size
                            except Exception as e:
                                print(f"Warning: Could not get image dimensions: {e}")
                                width, height = 400, 300  # default dimensions
                            
                            images_data.append({
                                "bytes": image_bytes,
                                "ext": image_format,
                                "width": width,
                                "height": height,
                                "paragraph_index": para_idx,
                                "paragraph_text": para_text,
                                "image_id": rel_id,
                                "source": "drawing"
                            })
                            print(f"Extracted drawing image {len(images_data)} from paragraph {para_idx}")
                        else:
                            print(f"Warning: Relationship ID {rel_id} not found in document parts")
                        
                except Exception as e:
                    print(f"Error extracting drawing image from DOCX paragraph {para_idx}: {e}")
                    continue
            
            # Look for pict elements (older format)
            for pict in run._element.xpath('.//w:pict'):
                try:
                    # Get image relationship ID from pict element
                    blip = pict.xpath('.//v:imagedata/@r:id', namespaces={
                        'v': 'urn:schemas-microsoft-com:vml',
                        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    })
                    
                    if blip:
                        rel_id = blip[0]
                        # Get the image from document relationships
                        if rel_id in doc.part.related_parts:
                            image_part = doc.part.related_parts[rel_id]
                            image_bytes = image_part.blob
                            
                            # Try to determine image format
                            image_format = 'png'  # default
                            if image_bytes.startswith(b'\xFF\xD8\xFF'):
                                image_format = 'jpeg'
                            elif image_bytes.startswith(b'\x89PNG'):
                                image_format = 'png'
                            elif image_bytes.startswith(b'GIF'):
                                image_format = 'gif'
                            elif image_bytes.startswith(b'BM'):
                                image_format = 'bmp'
                            
                            # Get image dimensions using PIL with better error handling
                            try:
                                with Image.open(BytesIO(image_bytes)) as img:
                                    width, height = img.size
                            except Exception as e:
                                print(f"Warning: Could not get image dimensions: {e}")
                                width, height = 400, 300  # default dimensions
                            
                            images_data.append({
                                "bytes": image_bytes,
                                "ext": image_format,
                                "width": width,
                                "height": height,
                                "paragraph_index": para_idx,
                                "paragraph_text": para_text,
                                "image_id": rel_id,
                                "source": "pict"
                            })
                            print(f"Extracted pict image {len(images_data)} from paragraph {para_idx}")
                        else:
                            print(f"Warning: Relationship ID {rel_id} not found in document parts")
                        
                except Exception as e:
                    print(f"Error extracting pict image from DOCX paragraph {para_idx}: {e}")
                    continue

    # Method 2: Extract from document parts directly (fallback method)
    if not images_data:
        print("No images found in paragraphs, trying document parts extraction...")
        try:
            for rel_id, part in doc.part.related_parts.items():
                if hasattr(part, 'blob') and part.content_type.startswith('image/'):
                    try:
                        image_bytes = part.blob
                        image_format = part.content_type.split('/')[-1]
                        
                        # Get image dimensions
                        try:
                            with Image.open(BytesIO(image_bytes)) as img:
                                width, height = img.size
                        except:
                            width, height = 400, 300
                        
                        images_data.append({
                            "bytes": image_bytes,
                            "ext": image_format,
                            "width": width,
                            "height": height,
                            "paragraph_index": 0,  # Default to first paragraph
                            "paragraph_text": "",
                            "image_id": rel_id,
                            "source": "document_parts"
                        })
                        print(f"Extracted image from document parts: {rel_id}")
                    except Exception as e:
                        print(f"Error extracting image from document parts {rel_id}: {e}")
        except Exception as e:
            print(f"Error in document parts extraction: {e}")

    print(f"Total extracted {len(images_data)} images from DOCX document")
    for i, img in enumerate(images_data):
        print(f"Image {i+1}: {img['source']} at paragraph {img['paragraph_index']}, format: {img['ext']}, size: {img['width']}x{img['height']}")
    
    return images_data

def associate_images_with_questions_docx(questions, images_data, doc):
    """Enhanced image association with questions using improved text-based matching"""
    if not images_data:
        print("No images found in DOCX document")
        return questions
        
    question_images = {}
    question_pattern = re.compile(r"Q(\d{1,9})\.")
    
    print(f"Starting image association for {len(questions)} questions and {len(images_data)} images")
    
    # Create a comprehensive text map of the document
    all_paragraphs = []
    question_locations = {}  # question_num -> paragraph_index
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        all_paragraphs.append((para_idx, text))
        
        # Track question boundaries
        match = question_pattern.match(text)
        if match:
            q_num = int(match.group(1))
            question_locations[q_num] = para_idx
            print(f"Found question Q{q_num} at paragraph {para_idx}: '{text[:50]}...'")
    
    # Determine question boundaries (where each question starts and ends)
    question_ranges = {}
    sorted_questions = sorted(question_locations.items())
    
    for i, (q_num, start_idx) in enumerate(sorted_questions):
        if i < len(sorted_questions) - 1:
            end_idx = sorted_questions[i + 1][1] - 1  # End before next question starts
        else:
            end_idx = len(all_paragraphs) - 1  # Last question goes to end
        question_ranges[q_num] = (start_idx, end_idx)
        print(f"Q{q_num} spans paragraphs {start_idx} to {end_idx}")
    
    # Associate each image with the question whose range it falls within
    for img_idx, img in enumerate(images_data):
        img_para = img["paragraph_index"]
        img_text = img.get("paragraph_text", "")
        
        print(f"Processing image {img_idx + 1} at paragraph {img_para} with text: '{img_text[:30]}...'")
        
        # Find which question range this image falls into
        associated_question = None
        
        for q_num, (start_idx, end_idx) in question_ranges.items():
            if start_idx <= img_para <= end_idx:
                associated_question = q_num
                print(f"Image {img_idx + 1} at paragraph {img_para} falls within Q{q_num} range ({start_idx}-{end_idx})")
                break
        
        # If not found in any range, find the closest preceding question
        if associated_question is None:
            closest_q = None
            min_distance = float('inf')
            
            for q_num, start_idx in question_locations.items():
                if start_idx <= img_para:
                    distance = img_para - start_idx
                    if distance < min_distance:
                        min_distance = distance
                        closest_q = q_num
            
            if closest_q:
                associated_question = closest_q
                print(f"Image {img_idx + 1} assigned to closest preceding question Q{closest_q}")
        
        # If still no association, assign to first question as fallback
        if associated_question is None and question_locations:
            associated_question = min(question_locations.keys())
            print(f"Image {img_idx + 1} assigned to first question Q{associated_question} as fallback")
        
        if associated_question:
            if associated_question not in question_images:
                question_images[associated_question] = []
            question_images[associated_question].append({
                "bytes": img["bytes"],
                "ext": img["ext"],
                "width": img["width"],
                "height": img["height"]
            })
            print(f"Image {img_idx + 1} successfully associated with Q{associated_question}")
    
    # Add images to questions
    updated_questions = []
    for question_text, existing_images in questions:
        # Extract question number
        match = question_pattern.match(question_text.strip())
        if match:
            q_num = int(match.group(1))
            question_imgs = question_images.get(q_num, [])
            if question_imgs:
                print(f"Adding {len(question_imgs)} images to Q{q_num}")
            # Combine existing images with newly found images
            all_images = existing_images + question_imgs
            updated_questions.append((question_text, all_images))
        else:
            updated_questions.append((question_text, existing_images))
    
    # Final summary
    total_with_images = sum(1 for _, imgs in updated_questions if imgs)
    print(f"Final result: {total_with_images} questions have images")
    for i, (q_text, imgs) in enumerate(updated_questions):
        if imgs:
            match = question_pattern.match(q_text.strip())
            q_num = match.group(1) if match else f"Q{i+1}"
            print(f"{q_num} has {len(imgs)} images")
    
    return updated_questions

def extract_questions_from_docx(file):
    doc = Document(file)
    questions = []
    current_question = ""
    question_pattern = re.compile(r"Q\d{1,9}\.")
    page_question_count = {0: 0}

    list_counter = {}

    print(f"Starting question extraction from DOCX with {len(doc.paragraphs)} paragraphs")

    # First pass: extract question text
    for para in doc.paragraphs:
        text = get_full_paragraph_text(para, list_counter)

        if not text:
            continue  # Skip empty paragraphs

        if question_pattern.match(text):
            if current_question:
                questions.append((current_question, []))  # Images will be added later
            current_question = text
            list_counter = {}  # Reset list counter for new question

        elif OPTION_LABEL_RE.match(text) or current_question:
            current_question += "\n" + text

    if current_question:
        questions.append((current_question, []))

    print(f"Extracted {len(questions)} questions from DOCX")

    # Extract images and associate with questions
    try:
        print("Starting image extraction and association...")
        images_data = extract_images_from_docx_paragraphs(doc)
        if images_data:
            print(f"Found {len(images_data)} images in DOCX")
            questions = associate_images_with_questions_docx(questions, images_data, doc)
        else:
            print("No images found in DOCX document")
    except Exception as e:
        print(f"Error processing DOCX images: {e}")
        import traceback
        traceback.print_exc()
        # Continue without images if there's an error

    return questions, page_question_count

def extract_images_with_positions_pdf(doc):
    """Enhanced PDF image extraction with better position tracking"""
    all_images = []
    
    print(f"Starting PDF image extraction from {len(doc)} pages")
    
    for page_number in range(len(doc)):
        page = doc.load_page(page_number)
        
        # Extract images with position info
        image_list = page.get_images(full=True)
        print(f"Found {len(image_list)} images on page {page_number + 1}")
        
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                
                # Get image rectangle (position on page)
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    img_rect = img_rects[0]  # Get first occurrence
                    
                    all_images.append({
                        "bytes": base_image["image"],
                        "ext": base_image["ext"],
                        "width": base_image["width"],
                        "height": base_image["height"],
                        "page": page_number,
                        "rect": img_rect,
                        "y_pos": img_rect.y0  # Top y-coordinate for sorting
                    })
                    print(f"Extracted image {img_index + 1} from page {page_number + 1}, position: {img_rect}")
                else:
                    print(f"Warning: No position found for image {img_index + 1} on page {page_number + 1}")
            except Exception as e:
                print(f"Error extracting image {img_index + 1} from PDF page {page_number + 1}: {e}")
                continue
    
    print(f"Total extracted {len(all_images)} images from PDF")
    return all_images

def associate_images_with_questions_pdf(questions_text, all_images, doc):
    """Enhanced PDF image-question association with better positioning logic"""
    question_images = {}
    question_pattern = re.compile(r"Q(\d{1,9})\.")
    
    print(f"Starting PDF image association for {len(questions_text)} questions and {len(all_images)} images")
    
    # Parse all pages to find question positions
    question_positions = {}
    
    for page_number in range(len(doc)):
        page = doc.load_page(page_number)
        text_blocks = page.get_text("dict")["blocks"]
        
        for block in text_blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        match = question_pattern.match(text)
                        if match:
                            q_num = int(match.group(1))
                            question_positions[q_num] = {
                                "page": page_number,
                                "y_pos": span["bbox"][1],  # Top y-coordinate
                                "bbox": span["bbox"]
                            }
                            print(f"Found Q{q_num} on page {page_number + 1} at y={span['bbox'][1]}")
    
    # Associate images with questions using improved logic
    for img_index, img in enumerate(all_images):
        img_page = img["page"]
        img_y = img["y_pos"]
        
        print(f"Processing image {img_index + 1} on page {img_page + 1} at y={img_y}")
        
        # Find questions on the same page
        page_questions = [(q_num, info) for q_num, info in question_positions.items() 
                         if info["page"] == img_page]
        
        if page_questions:
            # Sort questions by y-position (top to bottom, higher y values first in PDF coordinates)
            page_questions.sort(key=lambda x: x[1]["y_pos"], reverse=True)
            
            # Find the question that comes before this image (question above image)
            associated_question = None
            for q_num, q_info in page_questions:
                if q_info["y_pos"] >= img_y:  # Question is above or at same level as image
                    associated_question = q_num
                    break
            
            # If no question found above, associate with the first question on the page
            if not associated_question and page_questions:
                associated_question = page_questions[0][0]
                print(f"No question above image, associating with first question Q{associated_question}")
            
            if associated_question:
                if associated_question not in question_images:
                    question_images[associated_question] = []
                question_images[associated_question].append({
                    "bytes": img["bytes"],
                    "ext": img["ext"],
                    "width": img["width"],
                    "height": img["height"]
                })
                print(f"Image {img_index + 1} associated with Q{associated_question}")
        else:
            print(f"No questions found on page {img_page + 1} for image {img_index + 1}")
    
    # Add images to questions
    updated_questions = []
    for question_text in questions_text:
        # Extract question number
        match = question_pattern.match(question_text.strip())
        if match:
            q_num = int(match.group(1))
            question_imgs = question_images.get(q_num, [])
            if question_imgs:
                print(f"Adding {len(question_imgs)} images to Q{q_num}")
            updated_questions.append((question_text, question_imgs))
        else:
            updated_questions.append((question_text, []))
    
    # Final summary
    total_with_images = sum(1 for _, imgs in updated_questions if imgs)
    print(f"Final PDF result: {total_with_images} questions have images")
    
    return updated_questions

def extract_questions_from_pdf(pdf_data):
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    questions = []
    current_question = None
    question_pattern = re.compile(r"Q\d{1,9}\.")
    page_question_count = {}

    print(f"Starting PDF question extraction from {len(doc)} pages")

    # First pass: extract all question text
    questions_text = []
    for page_number in range(len(doc)):
        page = doc.load_page(page_number)
        text = page.get_text()
        
        # Count questions on this page
        question_matches = question_pattern.findall(text)
        page_question_count[page_number] = len(question_matches)
        print(f"Page {page_number + 1}: Found {len(question_matches)} questions")
        
        # Split text into lines for processing
        lines = text.split('\n')
        for line in lines:
            if question_pattern.match(line.strip()):
                # Save previous question if exists
                if current_question:
                    questions_text.append(current_question)
                
                # Start new question
                current_question = line
            elif current_question:
                # Accumulate lines for current question
                current_question += '\n' + line
    
    # Add last question if exists
    if current_question:
        questions_text.append(current_question)
    
    print(f"Extracted {len(questions_text)} questions from PDF text")
    
    # Extract all images with position information
    all_images = extract_images_with_positions_pdf(doc)
    
    # Associate images with questions
    questions = associate_images_with_questions_pdf(questions_text, all_images, doc)
    
    doc.close()
    return questions, page_question_count

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def force_table_indent_and_widths(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(Inches(0.2).pt)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    tbl.insert(0, tblPr)
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.85)

def add_image_to_docx_cell(cell, image_data):
    """Enhanced function to add image to DOCX table cell with better sizing and error handling"""
    try:
        # Create a new paragraph in the cell for the image
        img_paragraph = cell.add_paragraph()
        
        # Create image stream
        image_stream = BytesIO(image_data["bytes"])
        
        # Calculate appropriate size (max width 4 inches to fit in cell)
        max_width = Inches(4)
        max_height = Inches(3)
        
        original_width = image_data["width"]
        original_height = image_data["height"]
        aspect_ratio = original_height / original_width
        
        # Calculate dimensions maintaining aspect ratio
        if max_width * aspect_ratio <= max_height:
            img_width = max_width
            img_height = max_width * aspect_ratio
        else:
            img_height = max_height
            img_width = max_height / aspect_ratio
        
        # Add image to paragraph
        run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
        run.add_picture(image_stream, width=img_width, height=img_height)
        
        print(f"Successfully added image to DOCX cell: {img_width.inches:.2f}x{img_height.inches:.2f} inches")
        
    except Exception as e:
        print(f"Error adding image to DOCX: {e}")
        # Add error message instead of image
        error_para = cell.add_paragraph()
        error_para.add_run(f"[Image could not be loaded: {str(e)}]")

def process_question_block(block, positive, negative):
    block_text, images = block
    lines = [line for line in block_text.split("\n") if line.strip()]
    opts = []
    raw_options = []
    ans = ''
    sol_lines = []
    question_lines = []
    
    # Extract question number
    q_num = ""
    q_num_match = re.match(r"^(Q\d{1,9})\.", block_text.strip())
    if q_num_match:
        q_num = q_num_match.group(1)

    capturing_question = True
    capturing_option_index = -1
    capturing_solution = False

    for line in lines:
        if OPTION_LABEL_RE.match(line) and not capturing_solution:
            capturing_question = False
            capturing_solution = False
            raw_options.append(line)
            opts.append(line)  # Keep the full option line with label
            capturing_option_index = len(opts) - 1

        elif capturing_option_index != -1 and not line.lower().startswith(("correct answer", "solution")):
            # Append to current option with line break
            opts[capturing_option_index] += "\n" + line
            raw_options[-1] += "\n" + line

        elif line.lower().startswith("correct answer"):
            answer_text = line.split(":", 1)[-1].strip()
            ans_match = re.search(r"\b([A-Da-d1-4])\b", answer_text)
            if ans_match:
                ans_val = ans_match.group(1).upper()
                if ans_val in 'ABCD':
                    ans = str(ord(ans_val) - ord('A') + 1)
                else:
                    ans = ans_val
            capturing_option_index = -1
            capturing_solution = False

        elif line.lower().startswith("solution"):
            sol_lines.append(line.split(":", 1)[-1].strip())
            capturing_solution = True
            capturing_option_index = -1

        elif capturing_solution:
            sol_lines.append(line.strip())

        elif capturing_question:
            line = re.sub(r"^Q\d{1,9}\.\s*", "", line)
            question_lines.append(line)

    # Handle options - last 4 become the actual options
    if len(opts) > 4:
        # Keep original (labeled) versions of all extra options in question text
        question_lines.extend(raw_options[:-4])
        
        # For the final 4 options, strip the numbering
        final_options = [
            OPTION_LABEL_RE.sub("", opt.strip()).strip()
            for opt in raw_options[-4:]
        ]
    else:
        final_options = [
            OPTION_LABEL_RE.sub("", opt.strip()).strip()
            for opt in raw_options
        ] + [""] * (4 - len(raw_options))

    # Join question lines with line breaks
    q = "\n".join(question_lines)
    solution = "\n".join(sol_lines).strip()

    return {
        "Question": q.strip(),
        "Type": "multiple_choice",
        "Options": final_options,
        "Answer": ans,
        "Solution": solution,
        "Positive Marks": positive,
        "Negative Marks": negative,
        "Images": images,
        "Question Number": q_num
    }

def generate_docx(questions, bold_question=False):
    """Enhanced DOCX generation with better image handling"""
    document = Document()
    doc_stream = BytesIO()
    
    print(f"Generating DOCX for {len(questions)} questions")
    
    for question_index, data in enumerate(questions):
        print(f"Processing question {question_index + 1}: {data.get('Question Number', 'Unknown')}")
        
        table = document.add_table(rows=10, cols=2)
        table.autofit = False
        force_table_indent_and_widths(table)
        set_table_borders(table)

        labels = ["Question", "Type", "Option", "Option", "Option", "Option",
                "Answer", "Solution", "Positive Marks", "Negative Marks"]
        
        # Prepare values with proper line breaks
        values = [
            data["Question"], 
            data["Type"],
            *data["Options"][:4],  # Unpack the 4 options
            data["Answer"], 
            data["Solution"], 
            data["Positive Marks"], 
            data["Negative Marks"]
        ]

        for i, (label, value) in enumerate(zip(labels, values)):
            row = table.rows[i]
            row.cells[0].text = label
            
            # Clear existing content
            for paragraph in row.cells[1].paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            
            # Add new content with preserved formatting
            p = row.cells[1].add_paragraph()
            
            if label == "Question" and bold_question:
                run = p.add_run(value)
                run.bold = True
            elif label.startswith("Option"):
                # Add option text with original formatting
                for line in value.split('\n'):
                    p.add_run(line)
                # Remove last break if exists
                if p.runs and not p.runs[-1].text:
                    p._element.remove(p.runs[-1]._element)
            else:
                p.add_run(value)

            # Add images after question text
            if label == "Question" and data.get("Images"):
                print(f"Adding {len(data['Images'])} images to question {question_index + 1}")
                for img_idx, img_data in enumerate(data["Images"]):
                    print(f"Adding image {img_idx + 1} of {len(data['Images'])}")
                    add_image_to_docx_cell(row.cells[1], img_data)

        document.add_paragraph()
    
    print("Saving DOCX document...")
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

def generate_pdf(questions, bold_question=False):
    """Enhanced PDF generation with better image handling"""
    pdf_stream = BytesIO()
    doc = SimpleDocTemplate(pdf_stream, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    print(f"Generating PDF for {len(questions)} questions")

    if bold_question:
        bold_style = ParagraphStyle(
            'BoldQuestion',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            textColor=colors.red,
        )

    for question_index, data in enumerate(questions):
        print(f"Processing PDF question {question_index + 1}: {data.get('Question Number', 'Unknown')}")
        
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']

        def format_text_with_linebreaks(text):
            # Add <br/> before list-like patterns
            patterns = [
                r"\s([A-Za-z]\.)",
                r"\s(\d{1,2}\.)",
                r"\s([ivxlcdm]{1,4}\.)",
                r"\s([IVXLCDM]{1,4}\.)"
            ]

            for pattern in patterns:
                text = re.sub(pattern, r"<br/>&nbsp;\1", text)
            return text

        # Create question text with images
        question_elements = []
        question_elements.append(Paragraph(format_text_with_linebreaks(data["Question"]), bold_style if bold_question else normal_style))
        
        # Add images after question text
        if data.get("Images"):
            print(f"Adding {len(data['Images'])} images to PDF question {question_index + 1}")
            for img_idx, img_data in enumerate(data["Images"]):
                try:
                    print(f"Adding PDF image {img_idx + 1} of {len(data['Images'])}")
                    img_stream = BytesIO(img_data["bytes"])
                    
                    # Calculate image dimensions for PDF
                    original_width = img_data["width"]
                    original_height = img_data["height"]
                    aspect_ratio = original_height / original_width
                    
                    img_width = 4 * inch  # Max width to fit in table
                    img_height = img_width * aspect_ratio

                    if img_height > 2 * inch:  # Max height constraint
                        img_height = 2 * inch
                        img_width = img_height / aspect_ratio

                    question_elements.append(Spacer(1, 0.1 * inch))  # Space before image
                    question_elements.append(ReportLabImage(img_stream, width=img_width, height=img_height))
                    print(f"Successfully added PDF image: {img_width/inch:.2f}x{img_height/inch:.2f} inches")
                except Exception as e:
                    print(f"Error adding image to PDF: {e}")
                    question_elements.append(Paragraph(f"[Image could not be loaded: {str(e)}]", normal_style))

        # Create table data
        table_data = [
            ["Question", question_elements],
            ["Type", data["Type"]],
            ["Option A", Paragraph(data["Options"][0], normal_style)],
            ["Option B", Paragraph(data["Options"][1], normal_style)],
            ["Option C", Paragraph(data["Options"][2], normal_style)],
            ["Option D", Paragraph(data["Options"][3], normal_style)],
            ["Answer", data["Answer"]],
            ["Solution", Paragraph(format_text_with_linebreaks(data["Solution"]), normal_style)],
            ["Positive Marks", data["Positive Marks"]],
            ["Negative Marks", data["Negative Marks"]]
        ]

        # Create PDF table
        table = Table(table_data, colWidths=[1.5 * inch, 5 * inch])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))

        elements.append(table)
        elements.append(Spacer(1, 0.3 * inch))  # Space before next question

    # Build PDF
    print("Building PDF document...")
    doc.build(elements)
    pdf_stream.seek(0)
    return pdf_stream

@app.route('/upload', methods=['POST'])
def upload():
    file = request.files['pdf_file']
    filename = file.filename.lower()
    uploaded_data["original_filename"] = filename.rsplit('.', 1)[0]
    uploaded_data["positive"] = request.form.get('positive', '2')
    uploaded_data["negative"] = request.form.get('negative', '0.25')
    bold_question = request.form.get('bold_question', 'no') == 'yes'
    uploaded_data["bold_question"] = bold_question

    try:
        uploaded_data["range_start"] = int(request.form.get('range_start') or 1)
        uploaded_data["range_end"] = int(request.form.get('range_end') or 9999)
    except ValueError:
        return "❌ Invalid range input.", 400

    print(f"Processing file: {filename}")
    print(f"Question range: {uploaded_data['range_start']} to {uploaded_data['range_end']}")

    if filename.endswith(".pdf"):
        print("Extracting from PDF...")
        blocks, page_question_count = extract_questions_from_pdf(file.read())
    elif filename.endswith(".docx"):
        print("Extracting from DOCX...")
        blocks, page_question_count = extract_questions_from_docx(file)
    else:
        return "❌ Unsupported file format. Upload PDF or DOCX.", 400

    uploaded_data["blocks"] = blocks
    print(f"Extracted {len(blocks)} question blocks")

    errors = []
    base_numbers = []
    option_issues = []
    repeated_questions = []
    pattern = r"Q(\d{1,9})\."
    multi_page_warnings = []

    # Generate multi-page warnings
    for page_num, count in page_question_count.items():
        if count > 1:
            multi_page_warnings.append(f"Page {page_num+1} has {count} questions. Images on this page are associated with the first question that appears there.")

    for i, block in enumerate(blocks):
        block_text = block[0] if isinstance(block, tuple) else block
        match = re.match(pattern, block_text.strip())
        if match:
            num = int(match.group(1))
            base_numbers.append(num)
            
            # ✅ Sequence error check
            if i > 0 and base_numbers[i] != base_numbers[i-1] + 1:
                errors.append(f"Issue at Q{base_numbers[i]} (expected Q{base_numbers[i-1] + 1})")
        
            # ✅ Count options properly by line, not globally
            lines = block_text.strip().splitlines()
            # ✅ Remove redundant stripping - lines are already clean
            option_count = sum(1 for line in lines if OPTION_LABEL_RE.match(line))
            if option_count != 4:
                option_issues.append(f"Q{num} has {option_count} options")

    # ✅ Repeated questions
    counts = Counter(base_numbers)
    repeated_questions = [f"Q{num}" for num, count in counts.items() if count > 1]

    uploaded_data["base"] = base_numbers[0] if base_numbers else 1

    # ✅ Filter for selected question range
    filtered_qnums = []
    questions_to_generate = 0
    for block in blocks:
        block_text = block[0] if isinstance(block, tuple) else block
        match = re.match(pattern, block_text.strip())
        if match:
            q_num = int(match.group(1))
            if uploaded_data["range_start"] <= q_num <= uploaded_data["range_end"]:
                filtered_qnums.append(q_num)
                questions_to_generate += 1

    gen_start = min(filtered_qnums) if filtered_qnums else uploaded_data["range_start"]
    gen_end = max(filtered_qnums) if filtered_qnums else uploaded_data["range_end"]

    # ✅ Ensure lists are not None
    errors = errors or []
    option_issues = option_issues or []
    repeated_questions = repeated_questions or []
    multi_page_warnings = multi_page_warnings or []

    print(f"Diagnosis complete - Total questions: {len(blocks)}, Errors: {len(errors)}, Images found in document")

    return render_template("diagnose.html",
        total_qs=len(blocks),
        actual_start=base_numbers[0] if base_numbers else 0,
        actual_end=base_numbers[-1] if base_numbers else 0,
        range_start=uploaded_data["range_start"],
        range_end=uploaded_data["range_end"],
        base=uploaded_data["base"],
        option_issues=option_issues,
        errors=errors,
        repeated_questions=repeated_questions,
        questions_to_generate=questions_to_generate,
        gen_start=gen_start,
        gen_end=gen_end,
        multi_page_warnings=multi_page_warnings
    )

@app.route('/generate', methods=['POST'])
def generate():
    confirm = request.form.get("confirm", "no")
    output_format = request.form.get("format", "docx")
    blocks = uploaded_data["blocks"]
    positive = uploaded_data["positive"]
    negative = uploaded_data["negative"]
    range_start = uploaded_data["range_start"]
    range_end = uploaded_data["range_end"]
    bold_question = uploaded_data["bold_question"]  # Get the bold setting

    if confirm == "no":
        return redirect(url_for("index"))

    print(f"Generating {output_format.upper()} with range {range_start}-{range_end}")

    pattern = r"Q(\d{1,9})\."
    selected_blocks = []

    for block in blocks:
        block_text = block[0] if isinstance(block, tuple) else block
        match = re.match(pattern, block_text.strip())
        if match:
            q_num = int(match.group(1))
            if range_start <= q_num <= range_end:
                selected_blocks.append(block)

    if not selected_blocks:
        return "No questions found in the selected range.", 400

    print(f"Processing {len(selected_blocks)} selected questions")

    # Process all selected questions
    processed_questions = []
    for block_index, block in enumerate(selected_blocks):
        print(f"Processing block {block_index + 1}/{len(selected_blocks)}")
        data = process_question_block(block, positive, negative)
        processed_questions.append(data)
        
        # Log image information
        if data.get("Images"):
            print(f"Question {data.get('Question Number', block_index + 1)} has {len(data['Images'])} images")

    # Get a clean filename from the uploaded PDF name
    base_name = re.sub(r'[\\/*?:"<>|]', "_", uploaded_data.get("original_filename", "Processed_MCQs"))
    docx_filename = f"Bulk_Uploader_of_{base_name}.docx"
    pdf_filename = f"Bulk_Uploader_of_{base_name}.pdf"
    zip_filename = f"Bulk_Uploader_of_{base_name}.zip"

    # Handle different output formats
    if output_format == "docx":
        print("Generating DOCX output...")
        docx_stream = generate_docx(processed_questions, bold_question)
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=docx_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    elif output_format == "pdf":
        print("Generating PDF output...")
        pdf_stream = generate_pdf(processed_questions, bold_question)
        return send_file(
            pdf_stream,
            as_attachment=True,
            download_name=pdf_filename,
            mimetype="application/pdf"
        )

    elif output_format == "zip":
        print("Generating ZIP output with both DOCX and PDF...")
        # Create ZIP with both DOCX and PDF
        docx_stream = generate_docx(processed_questions, bold_question)
        pdf_stream = generate_pdf(processed_questions, bold_question)
        
        zip_stream = BytesIO()
        with ZipFile(zip_stream, 'w') as zipf:
            zipf.writestr(docx_filename, docx_stream.getvalue())
            zipf.writestr(pdf_filename, pdf_stream.getvalue())
        
        zip_stream.seek(0)
        return send_file(
            zip_stream,
            as_attachment=True,
            download_name=zip_filename,
            mimetype="application/zip"
        )

    return "❌ Only DOCX, PDF, and ZIP formats are supported on this server.", 400

if __name__ == "__main__":
    print("Starting Flask application...")
    app.run(host="0.0.0.0", debug=True)